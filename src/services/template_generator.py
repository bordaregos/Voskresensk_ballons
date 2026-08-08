"""Генератор шаблонов-заготовок .docx из ReportSchema.

Пишет документ через python-docx (не через сам docxtpl -- тот только
рендерит готовый .docx, не создаёт его). Результат -- РАЗОВАЯ заготовка:
дальше пользователь дорабатывает её в Word вручную (логотип, вёрстка,
формулировки), генератор этот файл больше не трогает -- safe regeneration
сознательно не реализовано (см. src/services/template_schema.py).

Ключевой инвариант всех add_*-функций: каждый {{ }}/{% %}-тег -- РОВНО ОДИН
run. python-docx даёт это бесплатно (paragraph.add_run(text) и cell.text=
всегда создают один run), если не разбивать один тег на несколько вызовов.
Это устраняет самый источник бага, найденного в реальном Шаблон_финал.docx:
там часть тегов раздроблена автозаменой Word на несколько <w:r>, из-за чего
их пришлось искать по документу через склейку рантайм-текста (см.
template_validator.py). Генератор такой проблемы не создаёт в принципе.

Второй подтверждённый на реальном шаблоне факт: цикл по списку -- обычный
Jinja `{% for x in list %}...{% endfor %}`, оборачивающий ЦЕЛИКОМ таблицу
(включая шапку) -- НЕ специальные докстпл-теги {%tr%}/{%p%} (их в реальном
шаблоне нет вообще). add_repeating_table() воспроизводит именно этот
паттерн.
"""

from pathlib import Path
from typing import Union

from docx import Document
from docx.document import Document as DocumentObject

from ..organization_config import OrganizationConfig
from .template_schema import (
    FieldsTableSection,
    ReportSchema,
    RepeatingTableSection,
    SCHEMAS,
    Section,
    StaticFieldsTableSection,
    StaticTextSection,
    TitleConfig,
)


def _add_heading(doc: DocumentObject, text: str, level: int = 2) -> None:
    doc.add_heading(text, level=level)


def _add_caption(doc: DocumentObject, caption: Union[str, None]) -> None:
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.italic = True


def add_title(doc: DocumentObject, title: TitleConfig) -> None:
    """Титульный лист: настраиваемый заголовок + плейсхолдеры отчёта."""
    heading = doc.add_heading(title.document_title, level=0)
    heading.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER
    for name in title.subtitle_fields:
        p = doc.add_paragraph()
        p.add_run("{{ " + name + " }}")


def add_static_text_section(doc: DocumentObject, section: StaticTextSection) -> None:
    """Раздел из чистого текста без плейсхолдеров (вводная часть и т.п.)."""
    if section.heading:
        _add_heading(doc, section.heading)
    for paragraph_text in section.paragraphs:
        doc.add_paragraph(paragraph_text)


def add_fields_table(doc: DocumentObject, section: FieldsTableSection) -> None:
    """Плоская таблица метка -> {{ placeholder }} (одна Jinja-переменная на
    строку, заполняется оператором через форму)."""
    if section.heading:
        _add_heading(doc, section.heading)
    _add_caption(doc, section.caption)
    table = doc.add_table(rows=len(section.rows), cols=2)
    table.style = "Table Grid"
    for row_idx, field_label in enumerate(section.rows):
        table.cell(row_idx, 0).text = field_label.label
        table.cell(row_idx, 1).text = "{{ " + field_label.placeholder + " }}"


def add_static_fields_table(
    doc: DocumentObject, section: StaticFieldsTableSection, org: OrganizationConfig,
) -> None:
    """Таблица реквизитов организации -- значения впечатываются буквально из
    org СЕЙЧАС, при генерации, а не оставляются Jinja-плейсхолдерами.
    Форма отчёта про эти поля ничего не знает, оператор их не заполняет --
    смешать эту таблицу с FieldsTableSection значило бы воссоздать тот же
    класс бага, ради устранения которого всё затевается (плейсхолдер,
    который никто и никогда не заполняет)."""
    if section.heading:
        _add_heading(doc, section.heading)
    _add_caption(doc, section.caption)
    table = doc.add_table(rows=len(section.rows), cols=2)
    table.style = "Table Grid"
    for row_idx, (label, attr_name) in enumerate(section.rows):
        table.cell(row_idx, 0).text = label
        table.cell(row_idx, 1).text = str(getattr(org, attr_name, ""))


def add_repeating_table(doc: DocumentObject, section: RepeatingTableSection) -> None:
    """{% for %}-параграф -> таблица (шапка + одна образцовая строка) ->
    {% endfor %}-параграф, строго последовательно. Вся таблица целиком
    повторяется по разу на элемент списка -- см. модульный докстринг."""
    if section.heading:
        _add_heading(doc, section.heading)
    _add_caption(doc, section.caption)

    for_paragraph = doc.add_paragraph()
    for_paragraph.add_run(
        "{% for " + section.loop_var + " in " + section.list_field + " %}"
    )

    table = doc.add_table(rows=2, cols=len(section.header_cells))
    table.style = "Table Grid"
    for col_idx, header_text in enumerate(section.header_cells):
        table.cell(0, col_idx).text = header_text

    if section.positional:
        for col_idx in range(len(section.header_cells)):
            table.cell(1, col_idx).text = f"{{{{ {section.loop_var}[{col_idx}] }}}}"
    else:
        for col_idx, attr_name in enumerate(section.row_cells):
            table.cell(1, col_idx).text = f"{{{{ {section.loop_var}.{attr_name} }}}}"

    endfor_paragraph = doc.add_paragraph()
    endfor_paragraph.add_run("{% endfor %}")


def add_section(doc: DocumentObject, section: Section, org: OrganizationConfig) -> None:
    if isinstance(section, StaticTextSection):
        add_static_text_section(doc, section)
    elif isinstance(section, StaticFieldsTableSection):
        add_static_fields_table(doc, section, org)
    elif isinstance(section, RepeatingTableSection):
        add_repeating_table(doc, section)
    elif isinstance(section, FieldsTableSection):
        add_fields_table(doc, section)
    else:
        raise TypeError(f"Неизвестный тип раздела схемы: {type(section)!r}")


def generate_template(
    equipment_type_id: str,
    org_config: OrganizationConfig,
    output_path: Union[str, Path],
    schema: Union[ReportSchema, None] = None,
    title_override: Union[str, None] = None,
) -> Path:
    """Собирает .docx-заготовку с нуля из ReportSchema + OrganizationConfig.

    Разовая операция: результат далее правится в Word вручную (логотип,
    точная вёрстка, формулировки), генератор его больше не трогает.
    title_override перекрывает schema.title.document_title без правки кода
    -- например, "сегодня отчёт, завтра заключение экспертизы".

    Raises:
        KeyError: неизвестный equipment_type_id и schema не передана явно.
    """
    if schema is None:
        schema = SCHEMAS[equipment_type_id]

    title = schema.title
    if title_override is not None:
        title = TitleConfig(document_title=title_override, subtitle_fields=title.subtitle_fields)

    doc = Document()
    add_title(doc, title)
    for section in schema.sections:
        add_section(doc, section, org_config)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
