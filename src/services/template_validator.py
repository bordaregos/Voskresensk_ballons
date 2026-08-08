"""Проверка .docx-шаблона на соответствие ожидаемому набору Jinja-полей.

Работает как со свежесгенерированным файлом (src/services/template_generator.py),
так и с уже отредактированным пользователем в Word -- это и есть основной
случай использования: шаблон генерируется один раз, дальше живёт своей
жизнью в Word, а валидатор даёт способ время от времени сверить его с
контрактом формы (src/ui/widget_names_*.py), не листая документ глазами.

Ключевой факт, на котором строится парсинг: python-docx уже сам склеивает
текст всех run параграфа в `paragraph.text`, независимо от того, на сколько
<w:r> Word разбил текст (а раздробление тегов автозаменой -- реальная
проблема, найденная в templates/Шаблон_финал.docx). Отдельной ручной склейки
XML не требуется -- extract_paragraph_text() ниже прямая обёртка над этим
фактом, названная отдельно ради ясности API и одного места, где это
задокументировано.
"""

import re
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from typing import Dict, List, Set, Union

from docx import Document

from .template_schema import FieldsTableSection, ReportSchema, RepeatingTableSection

# ``{{ field|filter }}`` -- Jinja-фильтр (например ``|lower``, см.
# final_condition в разделе 8.1) отбрасывается, полю нужно только имя.
_SCALAR_TAG_RE = re.compile(r"\{\{\s*([A-Za-z_][\w.]*)(?:\s*\|[^}]*)?\s*\}\}")
# ``{%tr for x in y %}``/``{%tc for x in y %}``/``{%p for x in y %}`` --
# докстпл-теги цикла по строке/ячейке/абзацу (без лишних пустых параграфов
# между итерациями, см. Таблица 2 и Таблица 4) -- то же самое, что обычный
# ``{% for x in y %}``, просто без пробела перед ``tr``/``tc``/``p``.
_FOR_TAG_RE = re.compile(r"\{%\s*(?:tr|tc|p)?\s*for\s+(\w+)\s+in\s+([\w.]+)\s*%\}")


class IssueLevel(Enum):
    ERROR = "error"      # ожидаемое поле схемы не найдено в документе
    WARNING = "warning"  # незнакомый плейсхолдер / подозрение на разрыв тега


@dataclass(frozen=True)
class ValidationIssue:
    level: IssueLevel
    message: str
    context: str = ""

    def __str__(self) -> str:
        prefix = "ОШИБКА" if self.level is IssueLevel.ERROR else "предупреждение"
        suffix = f" ({self.context})" if self.context else ""
        return f"[{prefix}] {self.message}{suffix}"


@dataclass(frozen=True)
class ValidationReport:
    issues: List[ValidationIssue] = field(default_factory=list)
    found_placeholders: Set[str] = field(default_factory=set)
    found_loop_fields: Set[str] = field(default_factory=set)

    @property
    def ok(self) -> bool:
        return not any(i.level is IssueLevel.ERROR for i in self.issues)

    @property
    def errors(self) -> List[ValidationIssue]:
        return [i for i in self.issues if i.level is IssueLevel.ERROR]

    @property
    def warnings(self) -> List[ValidationIssue]:
        return [i for i in self.issues if i.level is IssueLevel.WARNING]


def extract_paragraph_text(paragraph) -> str:
    """Текст параграфа независимо от того, на сколько run он раздроблен.

    python-docx (paragraph.text) уже делает это сам -- см. модульный
    докстринг. Отдельная функция нужна как явное, задокументированное
    место контракта, а не как самостоятельная логика склейки XML.
    """
    return paragraph.text


def _iter_table_cell_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                yield p
            for nested_table in cell.tables:
                yield from _iter_table_cell_paragraphs(nested_table)


def iter_document_paragraphs(doc):
    """Все параграфы документа: тело + ячейки всех таблиц (в т.ч. вложенных),
    без гарантии взаимного порядка между параграфами тела и таблиц (для
    задач валидации порядок не важен -- важен для генератора, см.
    tests/test_template_generator.py)."""
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        yield from _iter_table_cell_paragraphs(table)


def find_placeholders(doc) -> "tuple[Set[str], Set[str]]":
    """Возвращает (scalar_names, for_loop_field_names).

    Плейсхолдеры вида {{ var.attr }}, где var совпадает с именем переменной
    цикла из встреченного в документе {% for var in field %}, относятся к
    циклу field, а не считаются отдельным скалярным полем верхнего уровня
    -- иначе "{{ s.name }}" внутри цикла специалистов ошибочно считался бы
    отсутствующим полем "s.name" в self.data.
    """
    texts = [extract_paragraph_text(p) for p in iter_document_paragraphs(doc)]

    loop_vars: Dict[str, str] = {}
    loop_fields: Set[str] = set()
    for text in texts:
        for var, list_field in _FOR_TAG_RE.findall(text):
            loop_vars[var] = list_field
            loop_fields.add(list_field)

    scalars: Set[str] = set()
    for text in texts:
        for name in _SCALAR_TAG_RE.findall(text):
            prefix = name.split(".", 1)[0]
            if prefix in loop_vars:
                loop_fields.add(loop_vars[prefix])
            else:
                scalars.add(name)

    return scalars, loop_fields


def find_suspicious_tags(doc) -> List[ValidationIssue]:
    """'{{'/'{%'  без парного закрытия в пределах того же параграфа --
    признак непойманного разрыва тега (например, кто-то вручную начал
    печатать плейсхолдер в Word и автозамена/правка его разорвала)."""
    issues: List[ValidationIssue] = []
    for p in iter_document_paragraphs(doc):
        text = extract_paragraph_text(p)
        if not text.strip():
            continue
        open_curly = text.count("{{")
        close_curly = text.count("}}")
        open_pct = text.count("{%")
        close_pct = text.count("%}")
        if open_curly != close_curly or open_pct != close_pct:
            context = text if len(text) <= 120 else text[:117] + "..."
            issues.append(ValidationIssue(
                level=IssueLevel.WARNING,
                message="подозрение на раздроблённый/незакрытый Jinja-тег",
                context=context,
            ))
    return issues


def _expected_from_schema(schema: ReportSchema) -> "tuple[Set[str], Set[str]]":
    expected_scalars: Set[str] = set(schema.title.subtitle_fields)
    expected_loop_fields: Set[str] = set()
    for section in schema.sections:
        if isinstance(section, FieldsTableSection):
            expected_scalars.update(row.placeholder for row in section.rows)
        elif isinstance(section, RepeatingTableSection):
            expected_loop_fields.add(section.list_field)
    return expected_scalars, expected_loop_fields


def validate_template(
    docx_path: Union[str, Path], schema: ReportSchema, org_config=None,
) -> ValidationReport:
    """Сверяет .docx с ReportSchema. org_config принят для симметрии API,
    но не используется -- реквизиты организации (StaticFieldsTableSection)
    впечатываются буквально при генерации и намеренно не входят ни в
    ожидаемые, ни в проверяемые плейсхолдеры (см. template_schema.py)."""
    doc = Document(str(docx_path))
    found_scalars, found_loops = find_placeholders(doc)
    expected_scalars, expected_loops = _expected_from_schema(schema)

    issues: List[ValidationIssue] = []
    for name in sorted(expected_scalars - found_scalars):
        issues.append(ValidationIssue(
            IssueLevel.ERROR, f"поле «{name}» ожидается схемой, но не найдено в документе",
        ))
    for name in sorted(expected_loops - found_loops):
        issues.append(ValidationIssue(
            IssueLevel.ERROR,
            f"список «{name}» ожидается схемой, но не используется ни в одном {{% for %}}",
        ))
    for name in sorted(found_scalars - expected_scalars):
        issues.append(ValidationIssue(
            IssueLevel.WARNING, f"плейсхолдер «{{{{ {name} }}}}» не входит в схему (опечатка?)",
        ))
    for name in sorted(found_loops - expected_loops):
        issues.append(ValidationIssue(
            IssueLevel.WARNING, f"цикл по «{name}» не входит в схему (опечатка?)",
        ))
    issues.extend(find_suspicious_tags(doc))

    return ValidationReport(issues=issues, found_placeholders=found_scalars, found_loop_fields=found_loops)


def validate_against_widget_names(docx_path: Union[str, Path], widget_names_module) -> ValidationReport:
    """Упрощённый режим без ReportSchema -- прямая сверка с
    PLAIN_TEXT_EDIT_NAMES/COMBO_BOX_NAMES/DATE_EDIT_NAMES/SPIN_BOX_NAMES.
    Годится для проверки уже существующего Шаблон_финал.docx уже сегодня,
    не дожидаясь заполнения ReportSchema для баллонов.

    Список полей-циклов (ballony, bal_oval, tverdost_data, tables и т.п.) в
    widget_names_*.py не объявлен декларативно (собирается ad-hoc в
    main_window.py) -- в этом режиме такие поля не проверяются на
    отсутствие, только отражаются в found_loop_fields отчёта."""
    doc = Document(str(docx_path))
    found_scalars, found_loops = find_placeholders(doc)

    expected_scalars: Set[str] = set()
    for attr in ("PLAIN_TEXT_EDIT_NAMES", "COMBO_BOX_NAMES", "DATE_EDIT_NAMES", "SPIN_BOX_NAMES"):
        expected_scalars.update(getattr(widget_names_module, attr, []))

    issues: List[ValidationIssue] = []
    for name in sorted(expected_scalars - found_scalars):
        issues.append(ValidationIssue(
            IssueLevel.ERROR, f"поле «{name}» есть в widget_names, но не найдено в документе",
        ))
    for name in sorted(found_scalars - expected_scalars):
        issues.append(ValidationIssue(
            IssueLevel.WARNING, f"плейсхолдер «{{{{ {name} }}}}» не входит в widget_names (опечатка?)",
        ))
    issues.extend(find_suspicious_tags(doc))

    return ValidationReport(issues=issues, found_placeholders=found_scalars, found_loop_fields=found_loops)
