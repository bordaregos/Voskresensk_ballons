from docx import Document

from src.organization_config import DEFAULT_ORGANIZATION
from src.services.template_generator import generate_template
from src.services.template_schema import (
    FieldLabel,
    FieldsTableSection,
    RepeatingTableSection,
    ReportSchema,
    SCHEMAS,
    TitleConfig,
)
from src.services.template_validator import (
    IssueLevel,
    extract_paragraph_text,
    find_placeholders,
    find_suspicious_tags,
    validate_against_widget_names,
    validate_template,
)


def _doc_with_paragraph_runs(*run_texts):
    doc = Document()
    p = doc.add_paragraph()
    for text in run_texts:
        p.add_run(text)
    return doc, p


def test_extract_paragraph_text_joins_split_runs():
    doc, p = _doc_with_paragraph_runs("{{ na", "me", " }}")
    assert extract_paragraph_text(p) == "{{ name }}"


def test_find_placeholders_reconstructs_split_tag():
    doc, _ = _doc_with_paragraph_runs("{{ na", "me", " }}")
    scalars, loops = find_placeholders(doc)
    assert scalars == {"name"}
    assert loops == set()


def test_find_placeholders_extracts_loop_field_name():
    doc = Document()
    doc.add_paragraph("{% for j in bal_oval %}")
    doc.add_paragraph("{{ j.d_max_rand0 }}")
    doc.add_paragraph("{% endfor %}")
    scalars, loops = find_placeholders(doc)
    assert loops == {"bal_oval"}
    # j.d_max_rand0 относится к циклу, не должен попасть в скаляры верхнего уровня
    assert "j.d_max_rand0" not in scalars
    assert scalars == set()


def test_find_placeholders_recognizes_tr_for_loop_tag():
    """{%tr for %}/{%tr endfor %} -- докстпл-тег цикла по строке таблицы
    (без пустых параграфов между итерациями, см. Таблица 2/Таблица 4) --
    должен распознаваться так же, как обычный {% for %}."""
    doc = Document()
    doc.add_paragraph("{%tr for s in specialists %}")
    doc.add_paragraph("{{ s.position }}")
    doc.add_paragraph("{%tr endfor %}")
    scalars, loops = find_placeholders(doc)
    assert loops == {"specialists"}
    assert "s.position" not in scalars


def test_find_placeholders_strips_jinja_filter_from_scalar():
    """{{ final_condition|lower }} -- имя поля извлекается без фильтра."""
    doc = Document()
    doc.add_paragraph("{{ final_condition|lower }}")
    scalars, loops = find_placeholders(doc)
    assert scalars == {"final_condition"}


def test_find_suspicious_tags_detects_unclosed_tag():
    doc = Document()
    doc.add_paragraph("{{ broken")
    issues = find_suspicious_tags(doc)
    assert len(issues) == 1
    assert issues[0].level is IssueLevel.WARNING


def test_find_suspicious_tags_ignores_balanced_paragraph():
    doc = Document()
    doc.add_paragraph("{{ ok }} and {% for x in y %}")
    issues = find_suspicious_tags(doc)
    assert issues == []


def test_validate_against_own_generated_output_has_no_errors(tmp_path):
    path = generate_template("pipeline", DEFAULT_ORGANIZATION, tmp_path / "out.docx")
    report = validate_template(path, SCHEMAS["pipeline"])
    assert report.ok
    assert report.issues == []


def test_detects_missing_field_as_error(tmp_path):
    schema = ReportSchema(
        equipment_type_id="pipeline",
        title=TitleConfig(document_title="Т"),
        sections=[FieldsTableSection(heading=None, rows=[FieldLabel("Метка", "some_field")])],
    )
    doc = Document()
    doc.add_paragraph("документ без нужного плейсхолдера")
    path = tmp_path / "broken.docx"
    doc.save(str(path))

    report = validate_template(path, schema)
    assert not report.ok
    assert any("some_field" in i.message for i in report.errors)


def test_detects_missing_loop_field_as_error(tmp_path):
    schema = ReportSchema(
        equipment_type_id="pipeline",
        title=TitleConfig(document_title="Т"),
        sections=[RepeatingTableSection(
            heading=None, list_field="segments", loop_var="s",
            header_cells=["A"], row_cells=["a"],
        )],
    )
    doc = Document()
    doc.add_paragraph("нет цикла вообще")
    path = tmp_path / "broken.docx"
    doc.save(str(path))

    report = validate_template(path, schema)
    assert not report.ok
    assert any("segments" in i.message for i in report.errors)


def test_detects_unknown_placeholder_as_warning(tmp_path):
    schema = ReportSchema(
        equipment_type_id="pipeline",
        title=TitleConfig(document_title="Т"),
        sections=[FieldsTableSection(heading=None, rows=[FieldLabel("Метка", "known_field")])],
    )
    doc = Document()
    doc.add_paragraph("{{ known_field }}")
    doc.add_paragraph("{{ typo_field }}")
    path = tmp_path / "with_typo.docx"
    doc.save(str(path))

    report = validate_template(path, schema)
    assert report.ok  # опечатка -- предупреждение, не блокирует
    assert any("typo_field" in i.message for i in report.warnings)


def test_organization_fields_do_not_count_as_missing_or_extra(tmp_path):
    """Реквизиты организации (StaticFieldsTableSection) -- буквальный текст,
    не Jinja-плейсхолдеры. Валидатор не должен требовать их как поля схемы
    и не должен ругаться на них как на неизвестные плейсхолдеры (они и не
    появляются как {{ }}-теги в правильно сгенерированном документе)."""
    path_dir = tmp_path
    path = generate_template("pipeline", DEFAULT_ORGANIZATION, path_dir / "out.docx")
    report = validate_template(path, SCHEMAS["pipeline"])
    assert "full_name" not in report.found_placeholders
    assert not any("full_name" in i.message for i in report.issues)
    assert not any("address" in i.message for i in report.issues)


def test_validate_against_widget_names_catches_missing_and_unknown(tmp_path):
    class FakeWidgetNames:
        PLAIN_TEXT_EDIT_NAMES = ["report_number", "reg_number"]
        COMBO_BOX_NAMES = []
        DATE_EDIT_NAMES = []
        SPIN_BOX_NAMES = []

    doc = Document()
    doc.add_paragraph("{{ report_number }}")
    doc.add_paragraph("{{ unexpected_field }}")
    path = tmp_path / "partial.docx"
    doc.save(str(path))

    report = validate_against_widget_names(path, FakeWidgetNames)
    assert not report.ok
    assert any("reg_number" in i.message for i in report.errors)
    assert any("unexpected_field" in i.message for i in report.warnings)
