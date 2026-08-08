import re

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from src.organization_config import DEFAULT_ORGANIZATION
from src.services.template_generator import generate_template
from src.services.template_schema import (
    FieldsTableSection,
    RepeatingTableSection,
    ReportSchema,
    SCHEMAS,
    StaticFieldsTableSection,
)

TAG_RE = re.compile(r"\{\{.*?\}\}|\{%.*?%\}")


def _body_items(doc):
    """Элементы тела документа в порядке документа: ('p', Paragraph) или
    ('tbl', Table) -- doc.paragraphs/doc.tables по отдельности не сохраняют
    взаимный порядок, для проверки for/table/endfor нужен именно порядок."""
    items = []
    for child in doc.element.body:
        if child.tag == qn("w:p"):
            items.append(("p", Paragraph(child, doc)))
        elif child.tag == qn("w:tbl"):
            items.append(("tbl", Table(child, doc)))
    return items


def _all_paragraphs(doc):
    """Все параграфы документа: тело + все ячейки всех таблиц."""
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def test_generate_creates_valid_docx_file(tmp_path):
    path = generate_template("pipeline", DEFAULT_ORGANIZATION, tmp_path / "out.docx")
    assert path.exists()
    Document(str(path))  # не бросает исключение


def test_unknown_equipment_type_raises(tmp_path):
    with pytest.raises(KeyError):
        generate_template("unknown", DEFAULT_ORGANIZATION, tmp_path / "out.docx")


def test_generated_docx_contains_all_scalar_placeholders(tmp_path):
    path = generate_template("pipeline", DEFAULT_ORGANIZATION, tmp_path / "out.docx")
    doc = Document(str(path))
    full_text = "\n".join(p.text for p in _all_paragraphs(doc))

    expected = {
        field_label.placeholder
        for section in SCHEMAS["pipeline"].sections
        if isinstance(section, FieldsTableSection)
        for field_label in section.rows
    }
    assert expected, "схема должна содержать хотя бы одно скалярное поле"
    for placeholder in expected:
        assert "{{ " + placeholder + " }}" in full_text, placeholder


def test_repeating_table_wraps_whole_table_between_for_endfor(tmp_path):
    path = generate_template("pipeline", DEFAULT_ORGANIZATION, tmp_path / "out.docx")
    doc = Document(str(path))
    items = _body_items(doc)

    for_idx = next(
        i for i, (kind, obj) in enumerate(items)
        if kind == "p" and obj.text.strip() == "{% for s in specialists %}"
    )
    assert items[for_idx + 1][0] == "tbl"
    assert items[for_idx + 2][0] == "p"
    assert items[for_idx + 2][1].text.strip() == "{% endfor %}"


def test_repeating_table_sample_row_uses_loop_var_dot_attr(tmp_path):
    path = generate_template("pipeline", DEFAULT_ORGANIZATION, tmp_path / "out.docx")
    doc = Document(str(path))
    items = _body_items(doc)

    for_idx = next(
        i for i, (kind, obj) in enumerate(items)
        if kind == "p" and obj.text.strip() == "{% for s in specialists %}"
    )
    table = items[for_idx + 1][1]
    assert [c.text for c in table.rows[0].cells] == ["Должность", "ФИО", "Удостоверение"]
    assert [c.text for c in table.rows[1].cells] == [
        "{{ s.position }}", "{{ s.name }}", "{{ s.cert_number }}",
    ]


def test_repeating_table_positional_mode_uses_index(tmp_path):
    path = generate_template("pipeline", DEFAULT_ORGANIZATION, tmp_path / "out.docx")
    doc = Document(str(path))
    items = _body_items(doc)

    for_idx = next(
        i for i, (kind, obj) in enumerate(items)
        if kind == "p" and obj.text.strip() == "{% for doc in table_reviewed_docs %}"
    )
    table = items[for_idx + 1][1]
    assert [c.text for c in table.rows[1].cells] == ["{{ doc[0] }}", "{{ doc[1] }}"]


def test_each_tag_is_single_run(tmp_path):
    """Регрессия на баг из реального Шаблон_финал.docx: там часть тегов
    раздроблена автозаменой Word на несколько <w:r>. Сгенерированный
    документ обязан не воспроизводить эту проблему -- полный тег лежит в
    тексте ОДНОГО run, не только в объединённом paragraph.text."""
    path = generate_template("pipeline", DEFAULT_ORGANIZATION, tmp_path / "out.docx")
    doc = Document(str(path))

    checked_any = False
    for p in _all_paragraphs(doc):
        for tag in TAG_RE.findall(p.text):
            checked_any = True
            assert any(tag == run.text for run in p.runs), (
                f"тег {tag!r} раздроблен на несколько run: "
                f"{[r.text for r in p.runs]}"
            )
    assert checked_any, "в документе должен быть хотя бы один {{ }}/{% %}-тег"


def test_organization_fields_are_literal_not_placeholders(tmp_path):
    """StaticFieldsTableSection должна впечатывать значения OrganizationConfig
    буквально, а не оставлять {{ }}-теги -- иначе это ровно тот баг, ради
    устранения которого всё затевается. PIPELINE_SCHEMA этот тип секции
    больше не использует (1.2 стала обычными редактируемыми полями формы,
    org_* -- см. widget_names_pipeline.py), поэтому проверяем на синтетической
    схеме, как и test_custom_minimal_schema_round_trips."""
    from src.services.template_schema import ReportSchema, TitleConfig

    schema = ReportSchema(
        equipment_type_id="vessel",
        title=TitleConfig(document_title="Тестовый документ"),
        sections=[
            StaticFieldsTableSection(
                heading=None,
                rows=[("Экспертная организация", "full_name"), ("Адрес", "address")],
            ),
        ],
    )
    path = generate_template(
        "vessel", DEFAULT_ORGANIZATION, tmp_path / "out.docx", schema=schema,
    )
    doc = Document(str(path))
    full_text = "\n".join(p.text for p in _all_paragraphs(doc))

    assert DEFAULT_ORGANIZATION.full_name in full_text
    assert DEFAULT_ORGANIZATION.address in full_text
    # Ни один атрибут OrganizationConfig не должен всплыть как {{ }}-тег.
    for _label, attr_name in schema.sections[0].rows:
        assert "{{ " + attr_name + " }}" not in full_text


def test_title_override_replaces_document_title(tmp_path):
    path = generate_template(
        "pipeline", DEFAULT_ORGANIZATION, tmp_path / "out.docx",
        title_override="Заключение экспертизы промышленной безопасности",
    )
    doc = Document(str(path))
    assert doc.paragraphs[0].text == "Заключение экспертизы промышленной безопасности"


def test_custom_minimal_schema_round_trips(tmp_path):
    """Генератор работает на произвольной схеме, не только на встроенной --
    важно для будущих типов объектов (сосуды, котлы), не только трубопровода."""
    from src.services.template_schema import FieldLabel, TitleConfig

    schema = ReportSchema(
        equipment_type_id="vessel",
        title=TitleConfig(document_title="Тестовый документ"),
        sections=[
            FieldsTableSection(heading=None, rows=[FieldLabel("Поле", "custom_field")]),
        ],
    )
    path = generate_template(
        "vessel", DEFAULT_ORGANIZATION, tmp_path / "out.docx", schema=schema,
    )
    doc = Document(str(path))
    full_text = "\n".join(p.text for p in _all_paragraphs(doc))
    assert "{{ custom_field }}" in full_text
